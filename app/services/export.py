import logging
import re
from pathlib import Path

import docx
from docx.shared import Pt
import markdown
import weasyprint

logger = logging.getLogger(__name__)


def _add_rich_paragraph(doc, text: str, style=None):
    """Add a paragraph with bold markdown spans rendered as bold runs."""
    p = doc.add_paragraph(style=style)
    parts = re.split(r'(\*\*.+?\*\*)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)


def _add_table(doc, rows: list[list[str]]):
    """Add a markdown table as a DOCX table."""
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols, style="Light Grid Accent 1")
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            if j < n_cols:
                table.rows[i].cells[j].text = cell


def export_markdown(note_text: str, output_path: Path) -> Path:
    """Save note as Markdown file."""
    output_path = output_path.with_suffix(".md")
    output_path.write_text(note_text, encoding="utf-8")
    return output_path


def export_pdf(note_text: str, output_path: Path) -> Path:
    """Convert Markdown note to PDF."""
    output_path = output_path.with_suffix(".pdf")
    html_content = markdown.markdown(note_text, extensions=["extra", "sane_lists"])
    full_html = f"""<!DOCTYPE html>
<html lang="pl">
<head>
<meta charset="utf-8">
<style>
    body {{
        font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
        max-width: 800px;
        margin: 40px auto;
        padding: 0 20px;
        line-height: 1.6;
        color: #1a1a1a;
        font-size: 14px;
    }}
    h1, h2, h3 {{ color: #2c3e50; margin-top: 1.5em; }}
    h2 {{ border-bottom: 1px solid #eee; padding-bottom: 0.3em; }}
    ul {{ padding-left: 1.5em; }}
    li {{ margin-bottom: 0.3em; }}
    code {{ background: #f5f5f5; padding: 2px 6px; border-radius: 3px; }}
</style>
</head>
<body>{html_content}</body>
</html>"""
    weasyprint.HTML(string=full_html).write_pdf(str(output_path))
    return output_path


def export_docx(note_text: str, output_path: Path) -> Path:
    """Convert Markdown note to DOCX."""
    output_path = output_path.with_suffix(".docx")
    doc = docx.Document()

    table_rows = []
    for line in note_text.split("\n"):
        stripped = line.strip()

        # Collect markdown table rows
        if stripped.startswith("|"):
            # Skip separator rows like |---|---|
            if re.match(r'^\|[\s\-:|]+\|$', stripped):
                continue
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            table_rows.append(cells)
            continue

        # Flush collected table rows as a DOCX table
        if table_rows:
            _add_table(doc, table_rows)
            table_rows = []

        if not stripped:
            continue

        # Strip markdown bold markers for headings/paragraphs
        clean = re.sub(r'\*\*(.+?)\*\*', r'\1', stripped)

        if stripped.startswith("#### "):
            doc.add_heading(clean[5:], level=4)
        elif stripped.startswith("### "):
            doc.add_heading(clean[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(clean[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(clean[2:], level=1)
        elif stripped.startswith("- "):
            _add_rich_paragraph(doc, stripped[2:], style="List Bullet")
        elif stripped.startswith("* "):
            _add_rich_paragraph(doc, stripped[2:], style="List Bullet")
        elif re.match(r'^\d+\.\s', stripped):
            # Numbered list: "1. Item" -> List Number style
            text_after = re.sub(r'^\d+\.\s*', '', stripped)
            _add_rich_paragraph(doc, text_after, style="List Number")
        elif stripped.startswith("---"):
            continue  # Skip horizontal rules
        else:
            _add_rich_paragraph(doc, stripped)

    # Flush any remaining table rows
    if table_rows:
        _add_table(doc, table_rows)

    doc.save(str(output_path))
    return output_path


EXPORTERS = {
    "md": export_markdown,
    "pdf": export_pdf,
    "docx": export_docx,
}


def export_note(note_text: str, output_path: Path, fmt: str) -> Path:
    """Export note to the specified format."""
    exporter = EXPORTERS.get(fmt)
    if not exporter:
        raise ValueError(f"Unsupported export format: {fmt}. Use: {list(EXPORTERS.keys())}")
    result = exporter(note_text, output_path)
    logger.info("Exported note to %s (%d bytes)", result.name, result.stat().st_size)
    return result
